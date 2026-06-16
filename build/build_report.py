# -*- coding: utf-8 -*-
import datetime, io, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageChops

ROOT = r'C:\Users\Admin03\Desktop\proekt\mirigrushek'
OUT  = ROOT + r'\report.docx'
SCR  = ROOT + r'\build\_screens'

doc = Document()
st = doc.styles
st['Normal'].font.name = 'Arial'; st['Normal'].font.size = Pt(11)
for h, sz in [('Heading 1', 16), ('Heading 2', 13.5), ('Heading 3', 11.5)]:
    s = st[h]; s.font.name = 'Arial'; s.font.size = Pt(sz); s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)

sec = doc.sections[0]
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(sec, m, Inches(1))


def set_shading(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def code(text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(6)
    set_shading(p, 'F2F2F2')
    lines = text.split('\n')
    for i, line in enumerate(lines):
        r = p.add_run(line); r.font.name = 'Consolas'; r.font.size = Pt(size)
        if i < len(lines) - 1:
            r.add_break()
    return p


def para(text, bold=False, italic=False, size=11, align=None, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2)
        p.add_run(it)


def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        set_shading(t.rows[0].cells[i].paragraphs[0], 'F5DEB3')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(str(val))
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Inches(widths[i])
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_page_number(p):
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = 'PAGE'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(i); run._r.append(e)


def trim_bottom(path):
    im = Image.open(path).convert('RGB')
    bbox = ImageChops.difference(im, Image.new('RGB', im.size, (255, 255, 255))).getbbox()
    if bbox:
        im.crop((0, 0, im.width, min(im.height, bbox[3] + 14))).save(path)


def figure(path, caption, width=6.3):
    trim_bottom(path)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))
    para(caption, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)


fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run('Отчёт о работе над ИС «МирИгрушек» — стр. '); add_page_number(fp)

# ============================ Титульный лист ============================
for _ in range(3):
    doc.add_paragraph()
para('Демонстрационный экзамен', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('по специальности 09.02.07 «Информационные системы и программирование»', size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('Вариант В4 — профильный уровень (вариативная часть)', size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
doc.add_paragraph(); doc.add_paragraph()
para('ОТЧЁТ О ВЫПОЛНЕНИИ РАБОТЫ', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('информационная система магазина игрушек ООО «МирИгрушек»', bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
para('Описание выполненных действий и использованных команд', italic=True, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for _ in range(6):
    doc.add_paragraph()
para('Дата формирования: ' + datetime.date.today().strftime('%d.%m.%Y'), size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================ 1. Общие сведения ============================
doc.add_heading('1. Общие сведения о работе', level=1)
para('Разработана и развёрнута информационная система магазина игрушек ООО «МирИгрушек» — '
     'веб-приложение на стеке LAMP (Linux Ubuntu, веб-сервер Apache, СУБД MySQL, язык PHP). '
     'Реализованы: авторизация с разграничением прав по ролям, регистрация новых пользователей, '
     'каталог товаров, управление товарами и заказами. Ниже последовательно описаны выполненные '
     'действия и команды, которые я применял на каждом этапе.')
table(['Параметр', 'Значение'],
      [['Операционная система', 'Ubuntu Server 22.04 / 24.04 LTS'],
       ['Веб-сервер', 'Apache 2'],
       ['СУБД', 'MySQL 8'],
       ['Язык приложения', 'PHP 8'],
       ['Имя базы данных', 'mirigrushek (кодировка utf8mb4)'],
       ['Доступ к СУБД (MySQL Workbench)', 'host = <IP>, port = 3306, user = root, пароль 43215678$'],
       ['Вход в приложение', 'логин = фамилия латиницей @mail.ru, пароль 43215678$']],
      [2.6, 4.5])

# ============================ 2. Модуль 1: БД ============================
doc.add_heading('2. Модуль 1. Разработка базы данных средствами СУБД', level=1)
para('База данных приведена к третьей нормальной форме (3НФ): выделены справочники, устранены '
     'транзитивные зависимости, имена объектов заданы по индустриальным стандартам (snake_case). '
     'Состав заказа реализован связью «многие-ко-многим» через таблицу OrderItems. '
     'Полный SQL-код приведён в Приложении А.')
para('Перечень таблиц и связей:', bold=True, after=2)
table(['Таблица', 'Назначение', 'Внешние ключи'],
      [['Roles', 'Роли пользователей', '—'],
       ['Users', 'Пользователи', 'role_id → Roles'],
       ['Categories / Suppliers / Manufacturers / Units', 'Справочники товара', '—'],
       ['Products', 'Товары', 'unit_id, supplier_id, manufacturer_id, category_id'],
       ['PickupPoints', 'Пункты выдачи', '—'],
       ['OrderStatuses', 'Статусы заказов', '—'],
       ['Orders', 'Заказы', 'pickup_point_id, client_user_id, status_id'],
       ['OrderItems', 'Состав заказа', 'order_id → Orders, product_article → Products']],
      [3.0, 2.6, 3.5])

para('Вход в СУБД. Для работы с базой я подключался к MySQL командой:', after=2)
code('mysql -u root -p')
para('Создание базы данных. Создал базу и выбрал её для дальнейшей работы:', after=2)
code('CREATE DATABASE mirigrushek CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;\nUSE mirigrushek;')
para('Создание таблиц, полей, первичных ключей и связей (критерий 1.4). Пример создания таблицы '
     'товаров с первичным ключом, типами данных, ограничениями и внешним ключом:', after=2)
code('CREATE TABLE Products (\n'
     '  article         VARCHAR(20) PRIMARY KEY,\n'
     '  name            VARCHAR(500) NOT NULL,\n'
     '  unit_id         INT NOT NULL,\n'
     '  price           DECIMAL(10,2) NOT NULL CHECK (price >= 0),\n'
     '  supplier_id     INT NOT NULL,\n'
     '  manufacturer_id INT NOT NULL,\n'
     '  category_id     INT NOT NULL,\n'
     '  discount        INT NOT NULL DEFAULT 0,\n'
     '  stock_qty       INT NOT NULL DEFAULT 0,\n'
     '  description     TEXT,\n'
     '  photo           VARCHAR(255),\n'
     '  CONSTRAINT fk_prod_cat FOREIGN KEY (category_id) REFERENCES Categories(id)\n'
     ');')
para('Импорт данных (критерий 1.1). Данные из предоставленных файлов импорта я перенёс в виде '
     'команд INSERT и выполнил их. Пример наполнения таблицы товаров:', after=2)
code("INSERT INTO Products (article, name, unit_id, price, supplier_id, manufacturer_id,\n"
     "                      category_id, discount, stock_qty, description, photo) VALUES\n"
     "  ('PMEZMH', 'Набор машинок «Щенячий патруль»', 1, 1414, 1, 1, 1, 22, 50, '…', '1.jpg'),\n"
     "  ('BPV4MM', 'Конструктор «Сова Букля»',        1,  771, 2, 1, 2, 15, 26, '…', '2.jpg');")
para('Вывод данных. Чтобы убедиться в корректности загрузки, я выводил содержимое таблиц. '
     'Список всех таблиц базы я получил командой:', after=2)
code('SHOW TABLES;')
para('Структуру (поля и типы) конкретной таблицы я смотрел командой:', after=2)
code('DESCRIBE Products;')
para('Содержимое таблицы товаров я вывел запросом:', after=2)
code('SELECT * FROM Products;')
para('А, например, список пользователей с их ролями — запросом с объединением таблиц:', after=2)
code('SELECT u.full_name, u.login, r.name AS role\n'
     'FROM Users u JOIN Roles r ON r.id = u.role_id;')

# ============================ 3. Модуль 2 ============================
doc.add_heading('3. Модуль 2. Разработка приложения и вывод каталога', level=1)
para('Стартовый экран — окно входа. После авторизации в шапке отображается ФИО и роль '
     'пользователя. Каталог товаров выводится из базы данных карточками с фотографией '
     '(или заглушкой picture.png, если фото нет).')
para('Вывод каталога. Список товаров вместе со значениями из справочников я получаю запросом '
     'с объединением таблиц:', after=2)
code('SELECT p.*, c.name AS cat_name, s.name AS sup_name,\n'
     '       m.name AS man_name, u.name AS unit_name\n'
     'FROM Products p\n'
     '  JOIN Categories    c ON c.id = p.category_id\n'
     '  JOIN Suppliers     s ON s.id = p.supplier_id\n'
     '  JOIN Manufacturers m ON m.id = p.manufacturer_id\n'
     '  JOIN Units         u ON u.id = p.unit_id\n'
     'ORDER BY p.name;')
para('Подсветка строк (критерий 2.8) реализована по Руководству по стилю:', bold=True, after=2)
bullets([
    'если размер скидки превышает 17%, фон карточки товара — цвет #FFDEAD;',
    'если у товара снижена цена, основная цена перечёркнута и выделена красным, '
    'рядом указана итоговая цена чёрным цветом;',
    'если товара нет на складе, строка (карточка) выделяется голубым цветом;',
    'основной фон #FFFFFF, дополнительный #F5DEB3, акцент #DEB887; шрифт Arial; '
    'каждая форма имеет заголовок; на главной форме — логотип, у приложения — иконка.',
])

# ============================ 4. Модуль 3 ============================
doc.add_heading('4. Модуль 3. Последовательный интерфейс: добавление, '
                'редактирование, удаление товаров', level=1)
para('Реализованы формы добавления и редактирования товара (со всеми полями из задания и '
     'загрузкой фотографии с проверкой минимального размера 300×200 px), а также удаление '
     'товара с подтверждением. Ниже команды, которые выполняются при этих операциях.')
para('Добавление товара:', after=2)
code("INSERT INTO Products (article, name, unit_id, price, supplier_id, manufacturer_id,\n"
     "                      category_id, discount, stock_qty, description, photo)\n"
     "VALUES ('A1B2C3', 'Новый товар', 1, 990, 1, 1, 1, 0, 10, 'Описание', 'foto.jpg');")
para('Редактирование товара:', after=2)
code("UPDATE Products\n"
     "SET name = 'Новое название', price = 1290, discount = 20, stock_qty = 5\n"
     "WHERE article = 'A1B2C3';")
para('Удаление товара:', after=2)
code("DELETE FROM Products WHERE article = 'A1B2C3';")
para('Поиск по наименованию (критерий 3.7):', after=2)
code("SELECT * FROM Products WHERE name LIKE '%машин%';")
para('Сортировка по цене (критерий 3.8) и фильтрация по категории (критерий 3.9):', after=2)
code('SELECT * FROM Products WHERE category_id = 1 ORDER BY price ASC;')

# ============================ 5. Модуль 4 ============================
doc.add_heading('5. Модуль 4. Заказы: просмотр, добавление, редактирование, удаление', level=1)
para('По кнопке «Заказы» выводится список заказов согласно макету (артикулы, статус, адрес '
     'пункта выдачи, даты, клиент, код получения). Реализованы добавление, редактирование и '
     'удаление заказов.')
para('Вывод списка заказов с составом каждого заказа:', after=2)
code("SELECT o.id, o.order_date, o.delivery_date, os.name AS status,\n"
     "       pp.address, cu.full_name AS client, o.receive_code,\n"
     "       GROUP_CONCAT(CONCAT(oi.product_article,' ×',oi.quantity) SEPARATOR ', ') AS items\n"
     "FROM Orders o\n"
     "  LEFT JOIN PickupPoints  pp ON pp.id = o.pickup_point_id\n"
     "  LEFT JOIN Users         cu ON cu.id = o.client_user_id\n"
     "  JOIN      OrderStatuses os ON os.id = o.status_id\n"
     "  LEFT JOIN OrderItems    oi ON oi.order_id = o.id\n"
     "GROUP BY o.id ORDER BY o.id;")
para('Создание заказа и добавление его позиций:', after=2)
code("INSERT INTO Orders (order_date, delivery_date, pickup_point_id,\n"
     "                    client_user_id, receive_code, status_id)\n"
     "VALUES ('2025-05-01', '2025-05-05', 2, 7, '950', 2);\n"
     "INSERT INTO OrderItems (order_id, product_article, quantity)\n"
     "VALUES (LAST_INSERT_ID(), 'PMEZMH', 2);")
para('Изменение статуса заказа и удаление заказа:', after=2)
code("UPDATE Orders SET status_id = 1 WHERE id = 9;\n"
     "DELETE FROM Orders WHERE id = 9;")

# ============================ 6. Модуль 5: сервер ============================
doc.add_heading('6. Модуль 5. Развёртывание серверной инфраструктуры', level=1)
para('Ниже приведены команды, которые я последовательно выполнял на сервере (с правами '
     'администратора через sudo) для развёртывания системы.')

doc.add_heading('6.1 (5.1) Создание и конфигурирование виртуальной машины', level=2)
para('В среде виртуализации создал виртуальную машину с Ubuntu Server, настроил сеть (DHCP) '
     'и доступ по SSH. IP-адрес сервера определил командой:', after=2)
code('hostname -I')

doc.add_heading('6.2 (5.2) Установка и настройка серверной ОС', level=2)
para('Обновил списки пакетов системы:', after=2)
code('sudo apt update')
para('При необходимости обновил установленные пакеты:', after=2)
code('sudo apt upgrade -y')

doc.add_heading('6.3 (5.3) Установка и настройка СУБД', level=2)
para('Установил сервер MySQL:', after=2)
code('sudo apt install -y mysql-server')
para('Включил автозапуск и запустил службу:', after=2)
code('sudo systemctl enable --now mysql')
para('Чтобы к СУБД можно было подключаться по сети (из MySQL Workbench), в файле '
     '/etc/mysql/mysql.conf.d/mysqld.cnf задал параметр bind-address = 0.0.0.0 и перезапустил '
     'службу:', after=2)
code('sudo nano /etc/mysql/mysql.conf.d/mysqld.cnf\n'
     '# bind-address = 0.0.0.0\n'
     'sudo systemctl restart mysql')
para('Создал базу данных и таблицы, загрузил данные (см. разделы 2 и Приложение А). '
     'Настроил пользователя root для подключения с любого хоста и выдал ему все права:', after=2)
code("CREATE USER IF NOT EXISTS 'root'@'%' IDENTIFIED BY '43215678$';\n"
     "ALTER USER 'root'@'%' IDENTIFIED WITH caching_sha2_password BY '43215678$';\n"
     "GRANT ALL PRIVILEGES ON *.* TO 'root'@'%' WITH GRANT OPTION;\n"
     "FLUSH PRIVILEGES;")

doc.add_heading('6.4 (5.4) Установка и настройка веб-сервера', level=2)
para('Установил веб-сервер Apache, PHP и модуль связи PHP с Apache и MySQL:', after=2)
code('sudo apt install -y apache2 php libapache2-mod-php php-mysql php-mbstring')
para('Включил автозапуск и запустил Apache:', after=2)
code('sudo systemctl enable --now apache2')
para('Удалил стандартную приветственную страницу Apache, чтобы по адресу сервера сразу '
     'открывалось приложение:', after=2)
code('sudo rm /var/www/html/index.html')
para('Назначил index.php индексной страницей: в файле /etc/apache2/mods-enabled/dir.conf '
     'поставил index.php первым в списке DirectoryIndex, после чего перезапустил Apache:', after=2)
code('sudo nano /etc/apache2/mods-enabled/dir.conf\n'
     '# DirectoryIndex index.php index.html\n'
     'sudo systemctl reload apache2')

doc.add_heading('6.5 (5.5) Развёртывание информационной системы', level=2)
para('Скопировал файлы приложения в каталог веб-сервера и назначил владельцем пользователя '
     'веб-сервера:', after=2)
code('sudo cp -r web/*    /var/www/html/\n'
     'sudo cp -r images   /var/www/html/\n'
     'sudo chown -R www-data:www-data /var/www/html')
para('Открыл в брандмауэре порты 80 (веб-сайт) и 3306 (доступ к СУБД):', after=2)
code('sudo ufw allow 80/tcp\nsudo ufw allow 3306/tcp')
para('После этого сайт доступен в браузере по адресу http://<IP-адрес-сервера>/.')

# ============================ 7. Сводная таблица команд ============================
doc.add_heading('7. Сводная таблица использованных команд', level=1)
table(['Команда', 'Назначение'],
      [['hostname -I', 'Определение IP-адреса сервера'],
       ['sudo apt update / upgrade -y', 'Обновление списков и пакетов ОС'],
       ['sudo apt install -y mysql-server', 'Установка СУБД MySQL'],
       ['sudo apt install -y apache2 php …', 'Установка веб-сервера Apache и PHP'],
       ['sudo systemctl enable --now …', 'Автозапуск и запуск служб MySQL и Apache'],
       ['sudo systemctl restart / reload …', 'Перезапуск служб после изменения настроек'],
       ['mysql -u root -p', 'Подключение к СУБД'],
       ['CREATE DATABASE … / USE …', 'Создание и выбор базы данных'],
       ['CREATE TABLE …', 'Создание таблиц, полей, ключей и связей'],
       ['INSERT INTO …', 'Загрузка (импорт) данных в таблицы'],
       ['SHOW TABLES; / DESCRIBE …', 'Просмотр списка таблиц и структуры таблицы'],
       ['SELECT … (с JOIN, WHERE, ORDER BY)', 'Вывод и выборка данных, поиск, сортировка, фильтр'],
       ['UPDATE … / DELETE …', 'Редактирование и удаление товаров и заказов'],
       ['CREATE USER / GRANT / FLUSH', 'Настройка root для подключения отовсюду'],
       ['sudo rm /var/www/html/index.html', 'Удаление приветственной страницы Apache'],
       ['sudo cp -r … /var/www/html/', 'Размещение файлов приложения на сервере'],
       ['sudo chown -R www-data …', 'Назначение владельца файлов веб-сервера'],
       ['sudo ufw allow 80/3306', 'Открытие портов сайта и СУБД']],
      [3.7, 5.4])

# ============================ 8. Проверка ============================
doc.add_heading('8. Проверка работоспособности', level=1)
bullets([
    'Сайт открывается по адресу http://<IP-сервера>/ — отображается окно входа, затем каталог.',
    'Вход выполнен под ролями (администратор / менеджер / клиент), права доступа разграничены.',
    'Подключение в MySQL Workbench: host = <IP>, port = 3306, user = root, пароль 43215678$.',
])
para('Проведена отладка: приложение работает корректно, аварийного завершения не происходит '
     '(валидация ввода, обработка исключений, защита связанных записей от удаления). '
     'Скриншоты корректной работы приложения приведены ниже.', italic=True, after=8)
figure(SCR + r'\login.png',   'Рисунок 1 — окно входа и регистрации (стартовый экран)')
figure(SCR + r'\catalog.png', 'Рисунок 2 — каталог: фон #FFDEAD при скидке свыше 17%, '
       'перечёркнутая основная цена и итоговая цена, голубая подсветка товара, которого нет на складе')
figure(SCR + r'\orders.png',  'Рисунок 3 — список заказов (роль «Администратор»)')

# ============================ Приложение А: SQL ============================
doc.add_page_break()
doc.add_heading('Приложение А. Полный SQL-код базы данных', level=1)
para('Ниже приведён полный SQL-код, использованный для создания базы данных, всех таблиц, '
     'связей и загрузки данных.', after=8)
sql = io.open(ROOT + r'\init.sql', encoding='utf-8').read().strip('\n')
for block in re.split(r'\n\s*\n', sql):
    block = block.strip('\n')
    if block.strip():
        code(block, size=8)

doc.save(OUT)
print('saved', OUT)
